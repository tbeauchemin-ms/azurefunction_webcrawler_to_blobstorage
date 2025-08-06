import os
import re
import time
import uuid
import json
import requests
import hashlib
import threading
from collections import deque
from io import BytesIO
from urllib.parse import urlparse, urljoin, urlunparse
from urllib import robotparser
from threading import Lock
from xml.etree import ElementTree as ET
from readability import Document
from bs4 import BeautifulSoup
from azure.identity import DefaultAzureCredential, ManagedIdentityCredential
from azure.storage.blob import BlobServiceClient, ContentSettings
from playwright.sync_api import sync_playwright, TimeoutError as PlaywrightTimeoutError
from PyPDF2 import PdfReader
from docx import Document as DocxDocument
import pandas as pd
from requests.adapters import HTTPAdapter
from urllib3.util.retry import Retry
from concurrent.futures import ThreadPoolExecutor, wait, FIRST_COMPLETED
import openai
import tiktoken

# ---- LOGGING ----
def log(msg):
    pid = os.getpid()
    thread = threading.current_thread().name
    print(f"[PID {pid} | {thread}] {msg}", flush=True)

# ---- ENV/CONFIG ----
os_env = os.getenv

# ---- CRAWL CONFIGURATION ----
BASE_URLS = os_env("BASE_URLS", "")
ALLOW_DOMAINS = {d.strip().lower() for d in os_env("ALLOW_DOMAINS", "").split(";") if d.strip()}
SKIP_PATTERNS = [re.compile(p) for p in os_env("SKIP_REGEXES", "").split(";") if p]
MAX_WORKERS = int(os_env("MAX_WORKERS", "4"))
USER_AGENT = os_env("USER_AGENT", "Mozilla/5.0 (GenericCrawler/1.0)")
RESPECT_ROBOTS = os_env("RESPECT_ROBOTS", "true").lower() == "true"
SAVE_404 = os_env("SAVE_404", "false").lower() == "true"

# ---- NETWORK & TIMING CONFIGURATION ----
REQUEST_DELAY = float(os_env("REQUEST_DELAY", "0.5"))
PAGE_TIMEOUT_MS = int(os_env("PAGE_TIMEOUT_MS", "45000"))
NETWORK_IDLE_WAIT_MS = int(os_env("NETWORK_IDLE_WAIT_MS", "0"))
RETRY_COUNT = int(os_env("RETRY_COUNT", "3"))
RETRY_BACKOFF = int(os_env("RETRY_BACKOFF_FACTOR", "2"))

# ---- CONTENT PROCESSING CONFIGURATION ----
INCLUDE_PDFS = os_env("INCLUDE_PDFS", "true").lower() == "true"
INCLUDE_DOCX = os_env("INCLUDE_DOCX", "true").lower() == "true"
INCLUDE_XLSX = os_env("INCLUDE_XLSX", "true").lower() == "true"
MAX_CONTENT_CHARS = int(os_env("MAX_CONTENT_CHARS", "500000"))
MAX_DEPTH = int(os_env("MAX_DEPTH", "3"))
MAX_PDF_BYTES = int(os_env("MAX_PDF_BYTES", "20000000"))
PDF_CHUNK_CHARS = int(os_env("PDF_CHUNK_CHARS", "4000"))
PDF_CHUNK_OVERLAP = int(os_env("PDF_CHUNK_OVERLAP", "300"))
PDF_DOWNLOAD_TIMEOUT_MS = int(os_env("PDF_DOWNLOAD_TIMEOUT_MS", "60000"))

# ---- AZURE STORAGE CONFIGURATION ----
STORAGE_ACCOUNT = os.environ["STORAGE_ACCOUNT_NAME"]
STORAGE_AUTH_METHOD = os_env("STORAGE_AUTH_METHOD", "managedidentity")
STORAGE_CLIENT_ID = os_env("STORAGE_CLIENT_ID")
STORAGE_CONNECTION_STRING = os_env("STORAGE_CONNECTION_STRING")
CONTENT_CONT = os_env("CONTAINER_NAME", "content")
LOGS_CONT = os_env("LOG_CONTAINER_NAME", "logs")

# ---- AZURE OPENAI CONFIGURATION ----
OPENAI_ENDPOINT = os.environ["AZURE_OPENAI_ENDPOINT"]
OPENAI_API_KEY = os_env("AZURE_OPENAI_API_KEY")
OPENAI_API_VERSION = os_env("AZURE_OPENAI_API_VERSION", "2023-05-15")
OPENAI_AUTH_METHOD = os_env("AZURE_OPENAI_AUTH_METHOD", "managedidentity")
OPENAI_CLIENT_ID = os_env("AZURE_OPENAI_CLIENT_ID")
EMBEDDING_DEPLOYMENT_NAME = os.environ.get("AZURE_OPENAI_EMBEDDING_DEPLOYMENT_NAME", "text-embedding-ada-002") # This is the CUSTOM name YOU give the model (often the same as the actual).
EMBEDDING_MODEL_NAME = os.environ.get("AZURE_OPENAI_EMBEDDING_MODEL_NAME", "text-embedding-ada-002") # This is the actual model name.
EMBEDDING_TOKEN_LIMIT = int(os_env("EMBEDDING_TOKEN_LIMIT", "8191")) # Default to 8191 for ada-002 and 3-small
TOKEN_OVERLAP = int(os_env("TOKEN_OVERLAP", "300"))

# ---- MONITORING & LOGGING CONFIGURATION ----
INSIGHTS_KEY = os_env("APPINSIGHTS_INSTRUMENTATIONKEY", "")

# ----------- BlobServiceClient Auth Logic -----------
def create_blob_service_client():
    """
    Create Azure Blob Storage client based on the authentication method.
    Supports three methods: connectionstring, managedidentity (user-assigned), managedidentity (system-assigned)
    """
    if STORAGE_AUTH_METHOD.lower() == "connectionstring":
        if not STORAGE_CONNECTION_STRING:
            raise ValueError("STORAGE_CONNECTION_STRING must be set when using connectionstring authentication method")
        log(f"Creating Blob Storage client with connection string authentication")
        return BlobServiceClient.from_connection_string(STORAGE_CONNECTION_STRING)
    elif STORAGE_AUTH_METHOD.lower() == "managedidentity":
        log(f"Creating Blob Storage client with managed identity authentication")
        if STORAGE_CLIENT_ID:
            # Use user-assigned managed identity
            log(f"Using user-assigned managed identity for storage: {STORAGE_CLIENT_ID}")
            credential = ManagedIdentityCredential(client_id=STORAGE_CLIENT_ID)
        else:
            # Use system-assigned managed identity
            log(f"Using system-assigned managed identity for storage")
            credential = DefaultAzureCredential()

        return BlobServiceClient(
            account_url=f"https://{STORAGE_ACCOUNT}.blob.core.windows.net",
            credential=credential
        )
    else:
        raise ValueError(f"Unsupported storage authentication method: {STORAGE_AUTH_METHOD}. Supported methods: 'connectionstring', 'managedidentity'")

# ----------- OpenAI Client Auth Logic -----------
def create_openai_client():
    """
    Create Azure OpenAI client based on the authentication method.
    Supports three methods: apikey, managedidentity (user-assigned), managedidentity (system-assigned)
    """
    if OPENAI_AUTH_METHOD.lower() == "apikey":
        if not OPENAI_API_KEY:
            raise ValueError("AZURE_OPENAI_API_KEY must be set when using apikey authentication method")
        log(f"Creating OpenAI client with API key authentication")
        return openai.AzureOpenAI(
            api_key=OPENAI_API_KEY,
            api_version=OPENAI_API_VERSION,
            azure_endpoint=OPENAI_ENDPOINT
        )
    elif OPENAI_AUTH_METHOD.lower() == "managedidentity":
        log(f"Creating OpenAI client with managed identity authentication")
        if OPENAI_CLIENT_ID:
            # Use user-assigned managed identity
            log(f"Using user-assigned managed identity for OpenAI: {OPENAI_CLIENT_ID}")
            credential = ManagedIdentityCredential(client_id=OPENAI_CLIENT_ID)
        else:
            # Use system-assigned managed identity
            log(f"Using system-assigned managed identity for OpenAI")
            credential = DefaultAzureCredential()

        # Create a token provider function for Azure OpenAI
        def get_bearer_token_provider():
            token = credential.get_token("https://cognitiveservices.azure.com/.default")
            return token.token

        return openai.AzureOpenAI(
            azure_ad_token_provider=get_bearer_token_provider,
            api_version=OPENAI_API_VERSION,
            azure_endpoint=OPENAI_ENDPOINT
        )
    else:
        raise ValueError(f"Unsupported OpenAI authentication method: {OPENAI_AUTH_METHOD}. Supported methods: 'apikey', 'managedidentity'")

# Initialize blob service client
blob_service = create_blob_service_client()
content_container = blob_service.get_container_client(CONTENT_CONT)
logs_container = blob_service.get_container_client(LOGS_CONT)

visited = set()
documents = []
sitemap_urls = []
failed = []
robots_cache = {}
collection_lock = Lock()

# Dedup cache: (url, chunk_index)
seen = set()
try:
    with open("docs.json", "r") as f:
        for d in json.load(f):
            seen.add((d["url"], d.get("chunk_index", 1)))
except FileNotFoundError:
    pass

session = requests.Session()
session.headers.update({
    "User-Agent": USER_AGENT,
    "Accept": "application/pdf,*/*",
    "Accept-Language": "en-US,en;q=0.9",
    "Accept-Encoding": "gzip, deflate, br",
    "Connection": "keep-alive",
    "Referer": "https://www.google.com",
})
session.mount('https://', HTTPAdapter(max_retries=Retry(
    total=RETRY_COUNT, backoff_factor=RETRY_BACKOFF,
    status_forcelist=[403, 429, 500, 502, 503, 504]))
)

# ---- CORE UTILS, LOGGING EVERYWHERE ----
def get_embedding_aoai(text, client=None):
    enc = tiktoken.encoding_for_model(EMBEDDING_MODEL_NAME)
    num_tokens = len(enc.encode(text))
    if num_tokens > EMBEDDING_TOKEN_LIMIT:
        log(f"Skipping embedding for chunk (tokens={num_tokens} > {EMBEDDING_TOKEN_LIMIT})")
        return None
    if not text or not text.strip():
        log(f"Embedding skipped: empty input")
        return None

    if client is None:
        client = create_openai_client()
    try:
        response = client.embeddings.create(
            input=[text],
            model=EMBEDDING_DEPLOYMENT_NAME
        )
        log(f"Embedding generated for chunk ({len(text)} chars) using model {EMBEDDING_DEPLOYMENT_NAME} (api_version={OPENAI_API_VERSION})")
        return response.data[0].embedding
    except Exception as e:
        log(f"Embedding generation failed: {e}")
        return None


def get_domain_session(base_url: str) -> requests.Session:
    dom_sess = requests.Session()
    for prefix, adapter in session.adapters.items():
        dom_sess.mount(prefix, adapter)
    dom_sess.headers.update(session.headers)
    return dom_sess

def sanitize(text): return re.sub(r"[^A-Za-z0-9]+", "_", text).strip("_").lower()
def normalize_url(url):
    p = urlparse(url)
    scheme = "https"
    netloc = p.netloc.lower().replace("www.", "")
    path = p.path.rstrip("/")
    return urlunparse((scheme, netloc, path, "", "", ""))

def blob_name_for_url(url, chunk_index=None):
    parsed = urlparse(url)
    fragment = parsed.fragment or ""
    base_url = urlunparse((parsed.scheme, parsed.netloc.lower().replace("www.", ""), parsed.path.rstrip("/"), "", "", ""))
    url_hash = hashlib.md5((base_url + "#" + fragment).encode("utf-8")).hexdigest()
    base = sanitize(parsed.path or "root")
    if fragment:
        base = f"{base}_{sanitize(fragment)}"
    if chunk_index is not None:
        return f"{base}_{chunk_index}_{url_hash}.json"
    else:
        return f"{base}_{url_hash}.json"

def robot_allows(url, agent):
    if not RESPECT_ROBOTS:
        return True
    parsed = urlparse(url)
    base = f"{parsed.scheme}://{parsed.netloc}"
    if base not in robots_cache:
        rp = robotparser.RobotFileParser()
        rp.set_url(f"{base}/robots.txt")
        try:
            rp.read()
            robots_cache[base] = rp
        except:
            robots_cache[base] = None
    rp = robots_cache.get(base)
    return rp is None or rp.can_fetch(agent, url)

def should_skip(url):
    parsed = urlparse(url)
    if parsed.scheme not in ("http", "https"):
        return True
    for p in SKIP_PATTERNS:
        if p.search(url):
            return True
    return False

def is_allowed_link(link):
    host = urlparse(link).netloc.lower().replace("www.", "")
    return any(host == d or host.endswith(f".{d}") for d in ALLOW_DOMAINS)

def upload_json(data, blob_name, container, last_modified=None):
    try:
        meta = {}
        if last_modified:
            meta["last_modified"] = last_modified
        payload = json.dumps(data, ensure_ascii=False, indent=2).encode()
        cs = ContentSettings(content_type="application/json")
        blob_client = container.get_blob_client(blob_name)
        if blob_client.exists():
            props = blob_client.get_blob_properties()
            blob_lastmod = props.metadata.get("last_modified") if props.metadata else None
            if last_modified and blob_lastmod == last_modified:
                log(f"Skipping unchanged content: {blob_name}")
                return False
        blob_client.upload_blob(
            data=payload,
            overwrite=True,
            content_settings=cs,
            metadata=meta
        )
        log(f"Uploaded blob: {blob_name}")

        return True
    except Exception as e:
        log(f"Upload failed: {blob_name} - {e}")

        return False


def extract_main_content(html):
    """
    Use Readability to extract main article HTML, then strip out unwanted elements
    (headers, footers, sidebars, language‐selector blocks). Return cleaned text.
    """
    try:
        doc = Document(html)
        main_html = doc.summary()
        soup = BeautifulSoup(main_html, "html.parser")
    except Exception:
        soup = BeautifulSoup(html, "html.parser")

    # Remove common non‐content selectors
    for sel in [
        "[id*=lang]", "[class*=lang]", "[id*=translat]", "[class*=translat]",
        "[id*=goog]", "[class*=goog]", ".goog-te-banner-frame", ".goog-te-menu-frame",
        "[id*=cookie]", "[class*=cookie]", ".header", ".footer", ".sidebar",
        "header", "footer", "nav", "aside", "form", "script", "style"
    ]:
        for el in soup.select(sel):
            el.decompose()

    main = (
        soup.find("main") or
        soup.find(id="main-content") or
        soup.find("div", class_="content") or
        soup.body
    )
    text = main.get_text(separator="\n", strip=True) if main else soup.get_text(separator="\n", strip=True)

    # Filter out leftover language‐selector lines
    lines = text.splitlines()
    filtered = []
    in_lang_block = False
    for line in lines:
        if 'select language' in line.lower() or (len(line) > 40 and sum(1 for w in line.split() if len(w) > 5) > 4):
            in_lang_block = True
            continue
        if in_lang_block:
            if 'powered by translate' in line.lower() or not line.strip():
                in_lang_block = False
            continue
        filtered.append(line)

    cleaned_text = "\n".join(filtered).strip()
    return cleaned_text if len(cleaned_text) > 200 else text


def parse_sitemap(url, seen=None):
    """
    Recursively parse sitemap.xml. Return a list of all URLs under that sitemap, excluding ones that should_skip.
    """
    seen = seen or set()
    try:
        res = requests.get(url, headers={"User-Agent": USER_AGENT}, timeout=10)
        if res.status_code != 200:
            log(f"Sitemap fetch failed: {url}")
            return []
        root = ET.fromstring(res.content)
        ns = {"sm": "http://www.sitemaps.org/schemas/sitemap/0.9"}

        if root.tag.endswith("sitemapindex"):
            urls = []
            for loc in root.findall(".//sm:loc", ns):
                link = loc.text
                if link and link not in seen:
                    seen.add(link)
                    urls += parse_sitemap(link, seen)
            return urls

        return [
            loc.text for loc in root.findall(".//sm:loc", ns)
            if loc.text and not should_skip(loc.text)
        ]
    except Exception as e:
        log(f"Error parsing sitemap {url}: {e}")
        return []

# Regular expression to detect JS opens of file links
JS_OPEN_RE = re.compile(r'''javascript:window\.open\((["'])(.+?\.(pdf|docx|xlsx|doc|xls))\1''', re.IGNORECASE)


# ----------- Link Extraction and Normalization -----------
def extract_real_link(base_url, href):
    """
    If href is a JavaScript window.open to a known file extension, extract the URL.
    If href starts with '#', return base_url + that fragment (so #fragment remains).
    Otherwise, resolve href relative to base_url and normalize it.
    """
    if not href:
        return None

    href = href.strip()
    # 1) Handle javascript:window.open(...)
    if href.lower().startswith("javascript:"):
        m = JS_OPEN_RE.search(href)
        if m:
            relative_url = m.group(2)
            if relative_url.lower().endswith((".doc", ".xls")):
                return None
            return urljoin(base_url, relative_url)
        else:
            return None

    # 2) If href is just a "#fragment", keep the fragment
    if href.startswith("#"):
        return urljoin(base_url, href)

    # 3) Otherwise, resolve and normalize
    abs_url = urljoin(base_url, href)
    return normalize_url(abs_url)

def split_by_tokens(text, max_tokens=EMBEDDING_TOKEN_LIMIT-100, overlap=TOKEN_OVERLAP):
    enc = tiktoken.encoding_for_model(EMBEDDING_MODEL_NAME)
    tokens = enc.encode(text)
    chunks = []
    i = 0
    while i < len(tokens):
        chunk = tokens[i : i + max_tokens]
        chunk_text = enc.decode(chunk)
        chunks.append(chunk_text)
        i += max_tokens - overlap
    return chunks

# ----------- Chunking and Uploading -----------

def emit_chunks(url, text, last_mod):
    """
    Split text into token-length chunks (for Azure OpenAI embedding), upload as JSON docs.
    """
    windows = split_by_tokens(text)
    embedding_client = create_openai_client()
    for idx, chunk in enumerate(windows, 1):
        if not chunk or not chunk.strip():
            log(f"Skipping empty chunk {idx} for {url}")
            continue
        embedding = get_embedding_aoai(chunk, client=embedding_client)
        log(f"Chunk {idx} of {url}: embedding type: {type(embedding)}, first 5: {embedding[:5] if embedding else embedding}")
        rec = {
            "id": uuid.uuid4().hex,
            "url": url,
            "title": sanitize(urlparse(url).path),
            "chunk_index": idx,
            "chunk_total": len(windows),
            "content": chunk,
            "last_modified": last_mod,
            "embedding": embedding,
        }
        filename = blob_name_for_url(url, chunk_index=idx)
        key = (url, idx)
        if key not in seen and upload_json(rec, filename, content_container, last_modified=last_mod):
            with collection_lock:
                documents.append(rec)
                sitemap_urls.append(url)
            seen.add(key)


def handle_docx(url: str):
    norm_url = normalize_url(url)
    log(f"Visiting DOCX: {norm_url}")
    try:
        r = requests.get(url, headers={"User-Agent": USER_AGENT}, timeout=PDF_DOWNLOAD_TIMEOUT_MS / 1000)
        r.raise_for_status()
        doc = DocxDocument(BytesIO(r.content))
        content = "\n".join(para.text for para in doc.paragraphs)
        if not content.strip():
            return
        last_mod = r.headers.get("Last-Modified", time.strftime("%Y-%m-%dT%H:%M:%SZ", time.gmtime()))
        emit_chunks(norm_url, content, last_mod)
    except Exception as e:
        log(f"Error handling docx: {url} {e}")
        with collection_lock:
            failed.append({"url": url, "reason": f"DOCX parse error: {e}"})


def handle_xlsx(url: str):
    norm_url = normalize_url(url)
    log(f"Visiting XLSX: {norm_url}")
    try:
        r = requests.get(url, headers={"User-Agent": USER_AGENT}, timeout=PDF_DOWNLOAD_TIMEOUT_MS / 1000)
        r.raise_for_status()
        xls = pd.ExcelFile(BytesIO(r.content))
        text_chunks = []
        for sheet in xls.sheet_names:
            df = xls.parse(sheet, dtype=str, na_filter=False)
            rows = df.astype(str).values.tolist()
            for row in rows:
                text_chunks.append(" ".join(row))
        content = "\n".join(text_chunks)
        if not content.strip():
            return
        last_mod = r.headers.get("Last-Modified", time.strftime("%Y-%m-%dT%H:%M:%SZ", time.gmtime()))
        emit_chunks(norm_url, content, last_mod)
    except Exception as e:
        log(f"Error handling xlsx: {url} {e}")
        with collection_lock:
            failed.append({"url": url, "reason": f"XLSX parse error: {e}"})

def handle_pdf(playwright_ctx, url: str):
    norm_url = normalize_url(url)
    log(f"Visiting PDF: {norm_url}")
    parsed = urlparse(url)
    base_url = f"{parsed.scheme}://{parsed.netloc}"
    referer  = f"{base_url}/"
    pdf_sess = get_domain_session(base_url)
    pdf_sess.headers.update({
        "User-Agent": USER_AGENT,
        "Accept": "application/pdf,*/*;q=0.9",
        "Accept-Language": "en-US,en;q=0.9",
        "Accept-Encoding": "gzip, deflate, br",
        "Connection": "keep-alive",
        "Referer": referer,
        "Origin": base_url
    })
    retries = Retry(
        total=RETRY_COUNT,
        backoff_factor=RETRY_BACKOFF,
        status_forcelist=[403, 429, 500, 502, 503, 504],
        allowed_methods=["HEAD", "GET"]
    )
    pdf_sess.mount("https://", HTTPAdapter(max_retries=retries))
    pdf_sess.mount("http://", HTTPAdapter(max_retries=retries))
    try:
        pdf_sess.get(base_url, timeout=10).raise_for_status()
        h = pdf_sess.head(url, allow_redirects=True, timeout=10)
        if h.status_code == 200:
            size = int(h.headers.get("Content-Length", 0))
            if size > MAX_PDF_BYTES:
                log(f"Skipping large PDF: {norm_url} ({size} bytes)")
                return
        else:
            log(f"HEAD request returned {h.status_code} for {url}; proceeding to GET")
    except Exception as e:
        log(f"HEAD error for PDF {url}: {e}; proceeding to GET")
    success = False
    r = None  # Initialize r to avoid unbound variable error
    for attempt in range(RETRY_COUNT):
        try:
            r = pdf_sess.get(url, allow_redirects=True, timeout=(10, PDF_DOWNLOAD_TIMEOUT_MS / 1000))
            if r.status_code == 200 and "pdf" in r.headers.get("Content-Type", "").lower():
                success = True
                break
            else:
                with collection_lock:
                    failed.append({"url": url, "reason": f"PDF HTTP {r.status_code} (GET)"})
        except Exception as e:
            with collection_lock:
                failed.append({"url": url, "reason": f"PDF GET error: {e}"})
        time.sleep(REQUEST_DELAY * (RETRY_BACKOFF ** attempt))
    if success and r is not None:
        with collection_lock:
            failed[:] = [f for f in failed if not (f["url"] == url and "(GET)" in f["reason"])]
        reader = PdfReader(BytesIO(r.content))
        raw = "".join(p.extract_text() or "" for p in reader.pages)
        lm = r.headers.get("Last-Modified", time.strftime("%Y-%m-%dT%H:%M:%SZ", time.gmtime()))
        emit_chunks(norm_url, raw, lm)
        return
    log(f"All GET attempts failed; falling back to Playwright download")
    page = None  # Initialize page to avoid unbound variable error
    try:
        page = playwright_ctx.new_page()
        page.goto(base_url, timeout=PDF_DOWNLOAD_TIMEOUT_MS, wait_until="networkidle")
        with page.expect_download(timeout=PDF_DOWNLOAD_TIMEOUT_MS) as dl:
            page.goto(url, timeout=PDF_DOWNLOAD_TIMEOUT_MS)
        download = dl.value
        buffer = download.content()
        reader = PdfReader(BytesIO(buffer))
        raw = "".join(p.extract_text() or "" for p in reader.pages)
        lm = download.suggested_filename
        emit_chunks(norm_url, raw, lm)
        if page:
            page.close()
        return
    except Exception as e:
        log(f"Playwright browser download for {url} failed: {e}")
        with collection_lock:
            failed.append({"url": url, "reason": f"Playwright download error: {e}"})
        if page:
            try:
                page.close()
            except:
                pass
        return

def handle_page(playwright_ctx, url: str):
    url = enforce_trailing_slash_if_directory(url)
    norm_url = normalize_url(url)
    log(f"Visiting HTML: {norm_url}")
    if RESPECT_ROBOTS and not robot_allows(norm_url, USER_AGENT):
        return []
    pg = playwright_ctx.new_page()
    try:
        resp = None
        raw_links = []
        html = None
        soup = None
        headers = {}
        for attempt in range(RETRY_COUNT):
            try:
                resp = pg.goto(url, timeout=PAGE_TIMEOUT_MS, wait_until="networkidle")
                parsed = urlparse(url)
                fragment = parsed.fragment
                if fragment:
                    tab_selector = f"a[href='#{fragment}'], #{fragment}"
                    try:
                        pg.click(tab_selector, timeout=PAGE_TIMEOUT_MS)
                        if NETWORK_IDLE_WAIT_MS:
                            pg.wait_for_load_state("networkidle", timeout=NETWORK_IDLE_WAIT_MS)
                    except:
                        pass
                if resp.status == 429:
                    time.sleep(REQUEST_DELAY * (RETRY_BACKOFF ** attempt))
                    continue
                if resp.status >= 500:
                    raise PlaywrightTimeoutError(f"{resp.status}")
                html = pg.content()
                headers = resp.headers
                soup = BeautifulSoup(html, "html.parser")
                for nos in soup.select("noscript"):
                    nos.decompose()
                raw_links = pg.eval_on_selector_all(
                    "a[href]",
                    "els => els.map(e => e.getAttribute('href'))"
                )
                break
            except PlaywrightTimeoutError as e:
                log(f"HTML fetch attempt {attempt + 1} failed: {e}")
                time.sleep(REQUEST_DELAY * (RETRY_BACKOFF ** attempt))
            except Exception as e:
                log(f"HTML fetch attempt {attempt + 1} failed: {e}")
                time.sleep(REQUEST_DELAY * (RETRY_BACKOFF ** attempt))
        if soup is None:
            with collection_lock:
                failed.append({"url": url, "reason": "All fetch attempts failed or invalid HTML"})
            return []
        # Safely extract href attributes from anchor tags
        soup_links = []
        for a in soup.find_all("a", href=True):
            try:
                # Use getattr with default to safely access href
                href = getattr(a, 'get', lambda x, default=None: default)('href')
                if href:
                    soup_links.append(href)
            except (AttributeError, TypeError):
                continue
        all_raw_links = set(raw_links) | set(soup_links)
        log(f"Extracted {len(all_raw_links)} total <a> links from {url}")
        absolute_links = []
        for l in all_raw_links:
            real = extract_real_link(url, l)
            if real:
                absolute_links.append(real)
        filtered_links = [
            l for l in set(absolute_links)
            if is_allowed_link(l) and not should_skip(l)
        ]
        log(f"Filtered to {len(filtered_links)} crawlable links on {url}")
        if not resp:
            with collection_lock:
                failed.append({"url": url, "reason": "No response from server"})
            pg.close()
            return []
        if resp.status == 404:
            if SAVE_404:
                rec404 = {
                    "id": uuid.uuid4().hex,
                    "url": norm_url,
                    "content": "",
                    "status": 404
                }
                filename = f"404_{sanitize(urlparse(norm_url).path)}_{rec404['id']}.json"
                upload_json(rec404, filename, content_container)
            pg.close()
            return []
        if resp.status >= 400:
            with collection_lock:
                failed.append({"url": url, "reason": f"HTTP {resp.status}"})
            pg.close()
            return []
        fragments = set()
        for link in filtered_links:
            parsed = urlparse(link)
            if parsed.fragment:
                fragments.add(parsed.fragment)
        for fragment in fragments:
            pane = soup.find(id=fragment)
            if pane:
                pane_text = pane.get_text(separator="\n", strip=True)
                if pane_text.strip():
                    fragment_url = urljoin(url, f"#{fragment}")
                    last_mod = headers.get(
                        "Last-Modified",
                        time.strftime("%Y-%m-%dT%H:%M:%SZ", time.gmtime())
                    )
                    emit_chunks(fragment_url, pane_text, last_mod)
        rendered_html = str(soup)
        full_text = extract_main_content(rendered_html)
        if len(full_text) > MAX_CONTENT_CHARS:
            snippet = full_text[:MAX_CONTENT_CHARS]
            last = max(
                snippet.rfind("."),
                snippet.rfind("\n"),
                snippet.rfind(" ")
            )
            if last > 0:
                main_text = snippet[: last + 1].strip()
            else:
                main_text = snippet
        else:
            main_text = full_text
        last_mod = headers.get("Last-Modified", time.strftime("%Y-%m-%dT%H:%M:%SZ", time.gmtime()))
        emit_chunks(norm_url, main_text, last_mod)
        pg.close()
        return filtered_links
    except Exception as e:
        log(f"General error handling HTML {url}: {e}")
        with collection_lock:
            failed.append({"url": url, "reason": f"HTML error: {e}"})
        pg.close()
        return []


def enforce_trailing_slash_if_directory(url):
    parsed = urlparse(url)
    if "." not in os.path.basename(parsed.path):
        if not parsed.path.endswith("/"):
            return url.rstrip("/") + "/"
    return url

# ----------- BFS Crawler (Parallel) -----------

def crawl_worker(task):
    url, depth, max_depth = task
    norm_url = normalize_url(url)
    results = []
    log(f"WORKER starting {url} (depth={depth})")
    if depth > max_depth:
        return results
    if should_skip(norm_url) or (RESPECT_ROBOTS and not robot_allows(norm_url, USER_AGENT)):
        return results
    with sync_playwright() as pw:
        browser = pw.chromium.launch(
            headless=True,
            args=["--no-sandbox", "--disable-dev-shm-usage"]
        )
        context_opts = {
            "user_agent": USER_AGENT,
            "locale": "en-US",
            "accept_downloads": True,
            "java_script_enabled": True,
        }
        ctx = browser.new_context(**context_opts)
        try:
            lower_url = norm_url.lower()
            if lower_url.endswith(".pdf") and INCLUDE_PDFS:
                handle_pdf(ctx, url)
            elif lower_url.endswith(".docx") and INCLUDE_DOCX:
                handle_docx(url)
            elif lower_url.endswith(".xlsx") and INCLUDE_XLSX:
                handle_xlsx(url)
            elif lower_url.endswith(".doc") or lower_url.endswith(".xls"):
                log(f"Skipping unsupported Office format: {norm_url}")
            else:
                results = handle_page(ctx, url)
        except Exception as e:
            log(f"Worker error: {e}")
        finally:
            browser.close()
    return results


def bfs_crawl_parallel(seed_urls):
    global visited
    queue = deque()
    for seed_url in seed_urls:
        norm_seed = normalize_url(seed_url)
        if norm_seed not in visited:
            queue.append((seed_url, 0))
            visited.add(norm_seed)
    with ThreadPoolExecutor(max_workers=MAX_WORKERS) as executor:
        futures = dict()
        while queue or futures:
            while len(futures) < MAX_WORKERS and queue:
                url, depth = queue.popleft()
                task = (url, depth, MAX_DEPTH)
                log(f"Submitting {url} at depth {depth} to threadpool (queue size={len(queue)})")
                future = executor.submit(crawl_worker, task)
                futures[future] = (url, depth)
            done, _ = wait(futures, return_when=FIRST_COMPLETED)
            for future in done:
                url, depth = futures.pop(future)
                try:
                    links = future.result()
                    if links:
                        for link in links:
                            norm_link = normalize_url(link)
                            with collection_lock:
                                if norm_link not in visited:
                                    visited.add(norm_link)
                                    queue.append((link, depth + 1))
                                    log(f"Queued new URL: {link} (depth={depth+1})")
                except Exception as e:
                    log(f"Crawl task failed for {url}: {e}")

def start_crawl():
    log(f"MAX_WORKERS is {MAX_WORKERS}")
    try:
        seeds = [u.strip() for u in BASE_URLS.split(";") if u.strip()]
        for s in seeds:
            ALLOW_DOMAINS.add(urlparse(s).netloc.lower().replace("www.", ""))
        urls = []
        for seed in seeds:
            sm_url = seed.rstrip("/") + "/sitemap.xml"
            entries = parse_sitemap(sm_url)
            if entries:
                log(f"✅ Using sitemap for {seed} ({len(entries)} URLs)")
                for u in entries:
                    norm_u = normalize_url(u)
                    if norm_u not in visited:
                        urls.append(u)
            else:
                log(f"⚠️ No sitemap for {seed}; falling back to seed URL")
                urls.append(seed)
        log(f"🚀 Starting BFS crawl with {len(urls)} URLs (parallel)")
        bfs_crawl_parallel(urls)
    except Exception as e:
        log(f"Crawl failed: {e}")
    finally:
        if failed:
            upload_json(failed, "failed_urls.json", logs_container)
        if sitemap_urls:
            upload_json(sitemap_urls, "sitemap.json", logs_container)
        if documents:
            upload_json(documents, "docs.json", logs_container)
        log(f"Crawl complete. Total URLs Visited: {len(visited)}, Failed: {len(failed)}")

if __name__ == "__main__":
    start_crawl()
